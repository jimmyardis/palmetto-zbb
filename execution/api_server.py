#!/usr/bin/env python3
"""
Palmetto ZBB Suite — FastAPI Backend
=====================================
All dollar figures served from SQLite (integer cents) with source citations.
LLM used ONLY for reading and summarizing text — never for arithmetic.

Endpoints:
  GET  /health
  GET  /agencies
  GET  /agency/{section_number}
  POST /ask
  POST /scenario
  POST /sandbox/export
  GET  /                     → serves frontend
"""

import os
import io
import json
import math
import sqlite3
import logging
import hashlib
from pathlib import Path
from datetime import datetime, timezone
from typing import Optional

import anthropic
import voyageai
from pinecone import Pinecone
from fastapi import FastAPI, HTTPException, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field, field_validator
from dotenv import load_dotenv

# ─── Paths ──────────────────────────────────────────────────────────────────
ROOT    = Path(__file__).parent.parent
DB_PATH = ROOT / "budget_data.db"

# ─── Env ────────────────────────────────────────────────────────────────────
load_dotenv(ROOT.parent / ".env")
load_dotenv(ROOT / ".env")

ANTHROPIC_KEY  = os.environ.get("ANTHROPIC_API_KEY", "")
VOYAGE_KEY     = os.environ.get("VOYAGE_API_KEY", "")
PINECONE_KEY   = os.environ.get("PINECONE_API_KEY", "")
PINECONE_INDEX = os.environ.get("PINECONE_INDEX_BUDGET", "sc-budget")
PINECONE_NS    = "fy2026-zia"
VOYAGE_MODEL   = "voyage-3"
CLAUDE_MODEL   = "claude-sonnet-4-6"
MIN_SCORE      = 0.35
TOP_K          = 20
FISCAL_YEAR    = "2025-2026"

# ─── Logging ────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ─── Lazy clients ───────────────────────────────────────────────────────────
_voyage   = None
_pinecone = None
_claude   = None


def get_voyage():
    global _voyage
    if _voyage is None and VOYAGE_KEY:
        _voyage = voyageai.Client(api_key=VOYAGE_KEY)
    return _voyage


def get_pinecone():
    global _pinecone
    if _pinecone is None and PINECONE_KEY:
        pc = Pinecone(api_key=PINECONE_KEY)
        try:
            _pinecone = pc.Index(PINECONE_INDEX)
        except Exception as e:
            log.warning("Pinecone index not available: %s", e)
    return _pinecone


def get_claude():
    global _claude
    if _claude is None and ANTHROPIC_KEY:
        _claude = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    return _claude


# ─── DB connection ──────────────────────────────────────────────────────────
def get_db() -> sqlite3.Connection:
    if not DB_PATH.exists():
        raise HTTPException(500, "Budget database not found. Run Phase 1 ingestion first.")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA query_only = ON")
    return conn


def get_db_rw() -> sqlite3.Connection:
    if not DB_PATH.exists():
        raise HTTPException(500, "Budget database not found.")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


# ─── Money helpers ──────────────────────────────────────────────────────────
def cents_to_dollars_str(cents: int) -> str:
    """Format integer cents as dollar string for API responses."""
    sign = "-" if cents < 0 else ""
    return f"{sign}${abs(cents) // 100:,.0f}"


def apply_cut_integer(amount_cents: int, cut_pct: int) -> int:
    """
    Apply a percentage cut using integer arithmetic only.
    cut_pct: 0-100 (integer)
    Returns new amount in cents.
    No floats used.
    """
    if cut_pct <= 0:
        return amount_cents
    if cut_pct >= 100:
        return 0
    # Integer division: floor((amount * cut_pct) / 100)
    cut = (amount_cents * cut_pct) // 100
    return amount_cents - cut


# ─── FastAPI app ─────────────────────────────────────────────────────────────
CONFERENCE_NS = "fy2026-conference"

app = FastAPI(
    title="Palmetto ZBB Suite",
    description="Zero-based budgeting platform for the SC General Assembly. "
                "All dollar figures sourced verbatim from SQLite with traceable citations.",
    version="2.0.0",
)


# ════════════════════════════════════════════════════════════════════════════
# GET /health
# ════════════════════════════════════════════════════════════════════════════

@app.get("/health")
def health():
    """System status: DB row count, Pinecone vector count, reconciliation status."""
    status: dict = {
        "status": "ok",
        "fiscal_year": FISCAL_YEAR,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "database": {},
        "pinecone": {},
        "reconciliation": {},
    }

    # DB stats
    try:
        conn = get_db()
        row  = conn.execute("""
            SELECT
                COUNT(*) FILTER (WHERE is_total_row=0) AS data_rows,
                COUNT(DISTINCT section_number)          AS agencies,
                COUNT(*) FILTER (WHERE has_federal_match=1) AS federal_match_rows,
                COUNT(*) FILTER (WHERE extraction_confidence='low') AS low_conf
            FROM line_items
        """).fetchone()
        meta = {r["key"]: r["value"] for r in conn.execute("SELECT key,value FROM ingestion_meta").fetchall()}

        # Nonrecurring total from SQLite — no LLM arithmetic
        nr_row = conn.execute(
            "SELECT SUM(total_funds) AS s FROM line_items WHERE fund_category='nonrecurring'"
        ).fetchone()
        nonrecurring_cents = nr_row["s"] or 0

        recap_total_cents = int(meta.get("recap_total_cents", "0"))
        grand_total_cents = recap_total_cents + nonrecurring_cents

        status["database"] = {
            "data_rows": row["data_rows"],
            "agencies": row["agencies"],
            "federal_match_rows": row["federal_match_rows"],
            "low_confidence_rows": row["low_conf"],
            "last_ingestion": meta.get("last_reconciliation"),
            "db_path": str(DB_PATH),
        }
        status["reconciliation"] = {
            "status": meta.get("reconciliation_status", "unknown"),
            "recap_total": cents_to_dollars_str(recap_total_cents),
            "recap_gf": cents_to_dollars_str(int(meta.get("recap_gf_cents", "0"))),
            "recurring_total": cents_to_dollars_str(recap_total_cents),
            "nonrecurring_total": cents_to_dollars_str(nonrecurring_cents),
            "grand_total": cents_to_dollars_str(grand_total_cents),
            "recurring_total_cents": recap_total_cents,
            "nonrecurring_total_cents": nonrecurring_cents,
            "grand_total_cents": grand_total_cents,
        }
        conn.close()
    except Exception as e:
        status["database"] = {"error": str(e)}
        status["status"] = "degraded"

    # Pinecone stats
    try:
        idx = get_pinecone()
        if idx:
            info = idx.describe_index_stats()
            ns   = info.namespaces.get(PINECONE_NS, {})
            status["pinecone"] = {
                "index": PINECONE_INDEX,
                "namespace": PINECONE_NS,
                "vector_count": getattr(ns, "vector_count", 0),
                "available": True,
            }
        else:
            status["pinecone"] = {"available": False, "reason": "No API key or index not found"}
    except Exception as e:
        status["pinecone"] = {"available": False, "error": str(e)}

    return status


# ════════════════════════════════════════════════════════════════════════════
# GET /agencies
# ════════════════════════════════════════════════════════════════════════════

@app.get("/agencies")
def list_agencies(include_nonrecurring: bool = False):
    """
    All agency sections with totals from SQLite. No LLM involved.
    Dollar figures are verbatim from Part IA, FY2025-2026 (H.4025).
    ?include_nonrecurring=true appends the surplus and CRF fund entries.
    """
    conn = get_db()
    rows = conn.execute("""
        SELECT
            r.section_number,
            r.agency_name,
            r.total_funds       AS recap_total_cents,
            r.general_funds     AS recap_gf_cents,
            (r.total_funds - r.general_funds) AS recap_other_cents,
            COUNT(li.id)        AS line_item_count,
            SUM(li.has_federal_match) AS federal_match_count
        FROM recapitulation r
        LEFT JOIN line_items li
            ON li.section_number = r.section_number
            AND li.is_total_row = 0
        GROUP BY r.section_number, r.agency_name, r.total_funds, r.general_funds
        ORDER BY r.total_funds DESC
    """).fetchall()
    conn.close()

    agencies = []
    for r in rows:
        agencies.append({
            "section_number": r["section_number"],
            "agency_name": r["agency_name"],
            "total_funds_cents": r["recap_total_cents"],
            "total_funds_display": cents_to_dollars_str(r["recap_total_cents"] or 0),
            "general_funds_cents": r["recap_gf_cents"],
            "general_funds_display": cents_to_dollars_str(r["recap_gf_cents"] or 0),
            "other_funds_cents": r["recap_other_cents"],
            "other_funds_display": cents_to_dollars_str(r["recap_other_cents"] or 0),
            "line_item_count": r["line_item_count"],
            "federal_match_items": r["federal_match_count"],
            "citation": {
                "source_doc": "tap1a.htm",
                "section": f"Section {r['section_number']}",
                "fiscal_year": FISCAL_YEAR,
                "act": "H.4025, ratified May 28 2025",
            },
        })

    if include_nonrecurring:
        nr_rows = conn.execute("""
            SELECT line_item_description, total_funds, general_funds, source_doc, page_number
            FROM line_items
            WHERE fund_category = 'nonrecurring'
            ORDER BY page_number
        """).fetchall()
        conn.close()
        for nr in nr_rows:
            agencies.append({
                "section_number": None,
                "agency_name": nr["line_item_description"],
                "total_funds_cents": nr["total_funds"],
                "total_funds_display": cents_to_dollars_str(nr["total_funds"] or 0),
                "general_funds_cents": nr["general_funds"],
                "general_funds_display": cents_to_dollars_str(nr["general_funds"] or 0),
                "other_funds_cents": 0,
                "other_funds_display": "$0",
                "line_item_count": 1,
                "federal_match_items": 0,
                "fund_category": "nonrecurring",
                "citation": {
                    "source_doc": nr["source_doc"],
                    "page_number": nr["page_number"],
                    "fiscal_year": FISCAL_YEAR,
                    "act": "H.4025 + H.4026, ratified May 28 2025",
                },
            })
    else:
        conn.close()

    return {
        "fiscal_year": FISCAL_YEAR,
        "source": "H.4025, Part IA — SC FY2025-2026 Appropriations Act",
        "agency_count": len(agencies),
        "agencies": agencies,
    }


# ════════════════════════════════════════════════════════════════════════════
# GET /agency/{section_number}
# ════════════════════════════════════════════════════════════════════════════

@app.get("/agency/{section_number}")
def get_agency(section_number: str):
    """
    Full line item breakdown from SQLite + proviso text from Pinecone.
    All dollar figures from SQLite only. Citations on every figure.
    """
    conn = get_db()

    # Recap (authoritative agency total)
    recap = conn.execute(
        "SELECT * FROM recapitulation WHERE section_number=?", (section_number,)
    ).fetchone()
    if not recap:
        conn.close()
        raise HTTPException(404, f"Section {section_number} not found")

    # Line items
    items = conn.execute("""
        SELECT id, subsection_name, line_item_description,
               general_funds, federal_funds, other_funds, total_funds,
               source_doc, page_number, extraction_confidence,
               is_total_row, has_federal_match, federal_match_note
        FROM line_items
        WHERE section_number=? AND is_total_row=0
        ORDER BY id
    """, (section_number,)).fetchall()

    # Subtotals (is_total_row=1) — useful for frontend section grouping
    subtotals = conn.execute("""
        SELECT subsection_name, line_item_description,
               general_funds, other_funds, total_funds,
               source_doc, page_number
        FROM line_items
        WHERE section_number=? AND is_total_row=1
        ORDER BY id
    """, (section_number,)).fetchall()

    conn.close()

    line_items_out = []
    for it in items:
        line_items_out.append({
            "id": it["id"],
            "subsection": it["subsection_name"],
            "description": it["line_item_description"],
            "general_funds_cents": it["general_funds"],
            "general_funds_display": cents_to_dollars_str(it["general_funds"] or 0),
            "other_funds_cents": it["other_funds"],
            "other_funds_display": cents_to_dollars_str(it["other_funds"] or 0),
            "total_funds_cents": it["total_funds"],
            "total_funds_display": cents_to_dollars_str(it["total_funds"] or 0),
            "has_federal_match": bool(it["has_federal_match"]),
            "federal_match_note": it["federal_match_note"],
            "extraction_confidence": it["extraction_confidence"],
            "citation": {
                "source_doc": it["source_doc"],
                "page_number": it["page_number"],
                "section": f"Section {section_number}",
                "fiscal_year": FISCAL_YEAR,
                "act": "H.4025, ratified May 28 2025",
            },
        })

    # RAG: retrieve proviso text from Pinecone
    provisos = []
    try:
        voyage = get_voyage()
        idx    = get_pinecone()
        if voyage and idx:
            agency_name = recap["agency_name"] or ""
            query_text  = f"Section {section_number} {agency_name} budget provisions appropriations"
            emb = voyage.embed([query_text], model=VOYAGE_MODEL, input_type="query").embeddings[0]
            results = idx.query(
                vector=emb,
                top_k=TOP_K,
                namespace=PINECONE_NS,
                filter={"linked_section": {"$eq": section_number}},
                include_metadata=True,
            )
            # Also do a broader query without section filter
            if not results.matches or all(m.score < MIN_SCORE for m in results.matches):
                results = idx.query(
                    vector=emb,
                    top_k=TOP_K,
                    namespace=PINECONE_NS,
                    include_metadata=True,
                )

            for match in results.matches:
                if match.score >= MIN_SCORE:
                    m = match.metadata or {}
                    provisos.append({
                        "score": round(match.score, 4),
                        "text": m.get("text_preview", ""),
                        "source_doc": m.get("source_pdf", ""),
                        "page_number": m.get("page_number"),
                        "source_type": m.get("source_type", ""),
                        "linked_section": m.get("linked_section", ""),
                    })
    except Exception as e:
        log.warning("Pinecone query failed for section %s: %s", section_number, e)
        provisos = []

    return {
        "section_number": section_number,
        "agency_name": recap["agency_name"],
        "fiscal_year": FISCAL_YEAR,
        "totals": {
            "total_funds_cents": recap["total_funds"],
            "total_funds_display": cents_to_dollars_str(recap["total_funds"] or 0),
            "general_funds_cents": recap["general_funds"],
            "general_funds_display": cents_to_dollars_str(recap["general_funds"] or 0),
            "other_funds_cents": (recap["total_funds"] or 0) - (recap["general_funds"] or 0),
            "other_funds_display": cents_to_dollars_str(
                (recap["total_funds"] or 0) - (recap["general_funds"] or 0)
            ),
            "citation": {
                "source_doc": "tap1a.htm",
                "section": f"Recapitulation, Section {section_number}",
                "fiscal_year": FISCAL_YEAR,
                "act": "H.4025, ratified May 28 2025",
            },
        },
        "line_items": line_items_out,
        "line_item_count": len(line_items_out),
        "provisos": provisos,
        "data_note": (
            "Dollar figures are verbatim extractions from Part IA (tap1a.htm). "
            "other_funds = total − general (all non-general-fund appropriations). "
            "federal_funds not separately enumerated in Part IA source."
        ),
    }


# ════════════════════════════════════════════════════════════════════════════
# POST /ask
# ════════════════════════════════════════════════════════════════════════════

class AskRequest(BaseModel):
    question: str = Field(..., min_length=3, max_length=2000)
    section_filter: Optional[str] = None    # optional: limit to one section
    top_k: int = Field(default=20, ge=1, le=50)
    mode: str = Field(default="navigator")  # "navigator" | "suggest"


ANTI_HALLUCINATION_SYSTEM_PROMPT = """You are a neutral legislative budget analyst for the South Carolina General Assembly. You are assisting with analysis of the FY2025-2026 South Carolina Appropriations Act (H.4025, ratified May 28 2025).

CRITICAL RULES — YOU MUST FOLLOW THESE WITHOUT EXCEPTION:

1. You are strictly forbidden from stating any dollar amount, percentage, or fiscal figure that is not present verbatim in the context provided to you below.

2. If asked for a number you cannot find in the retrieved context, you MUST respond exactly: "I cannot confirm that figure from the source documents — please verify against the official appropriations act."

3. Never calculate, estimate, or approximate figures. Never derive figures by arithmetic from other figures unless the derived figure is explicitly stated in the context.

4. Always cite the specific section number, page number, and source document for every factual claim.

5. If the context does not contain sufficient information to answer the question, say so clearly rather than filling gaps with general knowledge.

6. You may summarize, explain, and contextualize the text in the context — but every specific number you mention must appear verbatim in the context.

7. Never speculate about legislative intent, future appropriations, or figures from prior years unless they appear in the provided context.

8. When a question asks for a list (e.g. "which agencies have X"), enumerate every instance found in the context. Do not stop at one example. If the context contains multiple matches, list all of them. If the context is likely incomplete for a comprehensive list, say so explicitly and recommend reviewing the full appropriations act.

The context below comes from official SC appropriations documents (Part IA appropriations tables and Part IB general provisions). Treat it as authoritative."""


@app.post("/ask")
def ask(req: AskRequest):
    """
    Plain-English RAG Q&A. LLM reads and summarizes text only — never does math.
    Every dollar figure in the response must appear verbatim in the retrieved context.
    """
    voyage = get_voyage()
    idx    = get_pinecone()
    claude = get_claude()

    if not voyage or not idx or not claude:
        raise HTTPException(503, "RAG services not available. Ensure VOYAGE_API_KEY, "
                                 "PINECONE_API_KEY, and ANTHROPIC_API_KEY are set and "
                                 "Phase 1C ingestion has been run.")

    # Embed query
    try:
        emb = voyage.embed([req.question], model=VOYAGE_MODEL, input_type="query").embeddings[0]
    except Exception as e:
        raise HTTPException(502, f"Embedding failed: {e}")

    # Retrieve from Pinecone — query both fy2026-zia (Part IB) and fy2026-conference
    query_kwargs: dict = {
        "vector": emb,
        "top_k": req.top_k,
        "namespace": PINECONE_NS,
        "include_metadata": True,
    }
    if req.section_filter:
        pass  # section_filter applied post-retrieval below (handles both list and string metadata formats)

    try:
        results = idx.query(**query_kwargs)
    except Exception as e:
        raise HTTPException(502, f"Vector retrieval failed: {e}")

    # Also query conference report namespace for adjustment justification context
    conf_matches = []
    try:
        conf_results = idx.query(
            vector=emb,
            top_k=max(2, req.top_k // 2),
            namespace=CONFERENCE_NS,
            include_metadata=True,
        )
        conf_matches = conf_results.matches or []
    except Exception:
        pass  # Conference namespace may not exist yet — degrade gracefully

    # Filter and deduplicate chunks
    chunks = []
    seen_ids = set()
    all_matches = list(results.matches) + conf_matches
    for match in all_matches:
        if match.score < MIN_SCORE:
            continue
        chunk_id = match.id
        if chunk_id in seen_ids:
            continue
        seen_ids.add(chunk_id)
        m = match.metadata or {}
        # Post-retrieval section filter. linked_section may be a list (new ingest) or
        # a comma-joined string (old ingest schema) — handle both.
        if req.section_filter:
            raw = m.get("linked_section", [])
            if isinstance(raw, list):
                linked = [str(s).strip() for s in raw if s]
            else:
                linked = [s.strip() for s in str(raw).split(",") if s.strip()]
            if req.section_filter not in linked:
                continue
            # In suggest mode, skip chunks that are only linked via General Provisions
            # (section 117) — those are cross-agency and rarely answer line-item-specific
            # questions. Only suppress if the chunk has no agency-specific link.
            if req.mode == "suggest" and linked == ["117"]:
                continue
        chunks.append({
            "score": match.score,
            "text": m.get("text_preview", ""),
            "source_pdf": m.get("source_pdf", ""),
            "page_number": m.get("page_number"),
            "linked_section": m.get("linked_section", ""),
            "source_type": m.get("source_type", ""),
            "description": m.get("description", ""),
        })

    # In suggest mode, if no chunk scores above 0.4 the context isn't useful — short-circuit.
    if req.mode == "suggest" and chunks and all(c["score"] < 0.40 for c in chunks):
        chunks = []

    if not chunks:
        if req.section_filter:
            if req.mode == "suggest":
                msg = f"No specific Part IB proviso applies to this line item in Section {req.section_filter}."
            else:
                msg = (
                    "No Part IB proviso text was found for this line item in the retrieved source documents. "
                    "This is common for allocation rows (ALLOC OTHER ENTITIES, pass-through funds) where "
                    "Part IB conditions are attached to the receiving agency's section rather than the "
                    "appropriating agency. Check the relevant program section in Part IB directly, or "
                    "consult the official appropriations act at "
                    "https://www.scstatehouse.gov/sess126_2025-2026/appropriations2025/tap1b.pdf"
                )
        else:
            msg = (
                "I cannot find relevant information in the retrieved source documents "
                "to answer this question. Please verify against the official appropriations act "
                "at https://www.scstatehouse.gov/sess126_2025-2026/appropriations2025/tap1a.htm"
            )
        return {
            "answer": msg,
            "citations": [],
            "chunks_retrieved": 0,
        }

    # Build context block
    context_parts = []
    for i, chunk in enumerate(chunks, 1):
        cite = f"[Source {i}: {chunk['source_pdf']}, Page {chunk['page_number']}"
        if chunk["linked_section"]:
            cite += f", Section {chunk['linked_section']}"
        cite += "]"
        context_parts.append(f"{cite}\n{chunk['text']}")

    context_text = "\n\n---\n\n".join(context_parts)

    # Generate answer
    if req.mode == "suggest":
        format_instruction = (
            "This is a ZBB justification suggestion request. "
            "If the context contains a specific Part IB proviso for this line item, "
            "quote the key requirement in 2-3 sentences and cite the source. "
            "If the context does not contain a directly applicable proviso, respond with exactly one sentence: "
            "'No specific Part IB proviso applies to this line item in Section {sec}.' "
            "Do not provide numbered lists, recommendations, links, or lengthy explanations."
        ).format(sec=req.section_filter or "this section")
    else:
        format_instruction = (
            "Answer using ONLY the information in the context above. "
            "Cite every factual claim with its source number."
        )

    user_message = (
        f"RETRIEVED CONTEXT FROM SC APPROPRIATIONS DOCUMENTS:\n\n"
        f"{context_text}\n\n"
        f"---\n\n"
        f"QUESTION: {req.question}\n\n"
        f"{format_instruction}"
    )

    try:
        response = claude.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1024,
            system=ANTI_HALLUCINATION_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_message}],
        )
        answer_text = response.content[0].text
    except Exception as e:
        raise HTTPException(502, f"LLM generation failed: {e}")

    # Build citations list
    citations = []
    for chunk in chunks:
        citations.append({
            "source_doc": chunk["source_pdf"],
            "page_number": chunk["page_number"],
            "section": chunk["linked_section"],
            "source_type": chunk["source_type"],
            "relevance_score": round(chunk["score"], 4),
            "text_preview": chunk["text"][:200],
            "fiscal_year": FISCAL_YEAR,
            "act": "H.4025, ratified May 28 2025",
        })

    return {
        "answer": answer_text,
        "citations": citations,
        "chunks_retrieved": len(chunks),
        "model": CLAUDE_MODEL,
        "system_prompt_active": "anti-hallucination v1",
    }


# ════════════════════════════════════════════════════════════════════════════
# GET /agency/{section_number}/insights
# ════════════════════════════════════════════════════════════════════════════

STRUCTURED_INSIGHTS_SYSTEM_PROMPT = """You are a senior ZBB analyst advising the South Carolina General Assembly.

Analyze the agency line items provided and respond in TWO parts — in this exact order:

PART 1 — A JSON block wrapped in <structured> tags. Produce one entry per program-area subsection present in the data. Each entry must have exactly these four keys:
- "subsection": the subsection name exactly as it appears in the data (e.g. "A. HOUSING, CARE, SECURITY, AND SUPERVISION")
- "recommended_tier": one of Mandated / High / Medium / Low
- "pre_fill_text": under 400 characters. Format: "Recommended tier: [TIER]. Before finalizing, verify: (1) ...; (2) ...; (3) ..." — specific, actionable questions referencing actual line item names and dollar amounts from the data.
- "peer_benchmark": under 200 characters. One concrete peer-state comparison labeled "Policy Analysis:" — or "Policy Analysis: No reliable benchmark available for this program type." if uncertain.

PART 2 — The full ZBB analyst report in markdown using this structure:
## Executive Summary
## Decision Unit Analysis
## Federal Match Risk Assessment
## Efficiency & Reform Opportunities
## Peer State Benchmarks (Policy Analysis)
## ZBB Analyst Action Checklist

Rules: SC dollar figures must match the data exactly. Peer benchmarks and estimates must be labeled "Policy Analysis". Do not fabricate peer figures.

Example output format:
<structured>
[
  {
    "subsection": "I. INTERNAL ADMIN & SUPPORT",
    "recommended_tier": "Medium",
    "pre_fill_text": "Recommended tier: Medium. Before finalizing, verify: (1) sub-object breakdown of Other Operating Expenses ($27,537,709); (2) FTE count vs. administrative FTE ratio; (3) shared-services consolidation opportunities with peer criminal justice agencies.",
    "peer_benchmark": "Policy Analysis: Administrative overhead in peer Southern DOC agencies typically runs 5-8% of total budget. SC Internal Admin is approximately 6.5% — within normal range."
  }
]
</structured>

## Executive Summary
[narrative continues here...]"""


INSIGHTS_SYSTEM_PROMPT = """You are a senior ZBB (zero-based budgeting) analyst advising the South Carolina General Assembly.

Your task: analyze the actual H.4025 line-item appropriations for a specific SC state agency and produce a structured, actionable ZBB report that legislative staff and fiscal analysts can use to fill in Justified Amounts and write justifications in the Decision Package.

YOU HAVE TWO TYPES OF KNOWLEDGE:

1. SC BUDGET FACTS (line-item data provided): Authoritative. When citing a specific dollar amount from the data, attribute it to H.4025. Never fabricate or calculate dollar figures not in the data.

2. POLICY ANALYSIS & PEER BENCHMARKS (from your training): You ARE permitted — and expected — to draw on your knowledge of corrections systems, privatization practices, per-inmate costs in peer states, federal grant programs, staffing benchmarks, and ZBB reform best practices. Clearly label this content as "Policy Analysis" or "Peer Benchmark" so readers know it is not a SC budget figure.

OUTPUT FORMAT — produce all of the following sections in markdown:

## Executive Summary
3-4 sentences: agency budget profile, the dominant cost driver, General Fund dependency, and the core ZBB opportunity.

## Decision Unit Analysis
For each program area / subsection, include:
- Program name and total appropriation (from data, cite H.4025)
- Recommended ZBB Priority Tier: Mandated / High / Medium / Low — with a 1-sentence rationale
- 2-3 analyst questions that must be answered before the Justified Amount can be set
- Any efficiency or privatization flags specific to this line

## Federal Match Risk Assessment
For any lines flagged with federal match exposure: explain what happens to federal dollars if the GF line is reduced, and what the analyst needs to verify.

## Efficiency & Reform Opportunities
For each opportunity:
- Specific line item(s) involved (with H.4025 dollar amount)
- Reform approach (privatization, consolidation, outsourcing, etc.)
- What peer states do (FL, GA, NC, TN, TX) — label as Policy Analysis
- Realistic savings range — label clearly as estimate, not SC data

## Peer State Benchmarks (Policy Analysis)
2-3 concrete comparisons to peer Southern states on corrections-relevant metrics (cost per inmate, staffing ratios, privatization share, etc.). Label all as Policy Analysis.

## ZBB Analyst Action Checklist
A numbered list of the 8-12 most important verification steps before the Decision Package is finalized for this agency.

Be specific — reference actual line items by name and dollar amount. Be honest about uncertainty. Do not fabricate peer state figures."""


def _get_structured_insights(section_number: str, agency_name: str, items) -> dict:
    """
    Internal helper: calls Claude with STRUCTURED_INSIGHTS_SYSTEM_PROMPT and returns
    {"units": [...], "narrative": "..."}.  Falls back to {"units": [], "narrative": ""} on any error.
    """
    from collections import defaultdict
    import json as _json

    claude = get_claude()
    if not claude:
        return {"units": [], "narrative": ""}

    subsections: dict = defaultdict(list)
    for it in items:
        subsections[it["subsection_name"] or "GENERAL"].append(it)

    prompt_lines = [
        f"AGENCY: {agency_name}",
        f"SECTION: {section_number}",
        f"FISCAL YEAR: {FISCAL_YEAR}",
        f"SOURCE: H.4025, SC FY{FISCAL_YEAR} Appropriations Act, ratified May 28 2025",
        "",
        "LINE ITEMS BY PROGRAM AREA:",
    ]
    for subsec, sub_items in subsections.items():
        prompt_lines.append(f"\n--- {subsec} ---")
        for it in sub_items:
            gf  = cents_to_dollars_str(it["general_funds"] or 0)
            oth = cents_to_dollars_str(it["other_funds"] or 0)
            tot = cents_to_dollars_str(it["total_funds"] or 0)
            fed = "  ⚠ FED MATCH" if it["has_federal_match"] else ""
            prompt_lines.append(f"  {it['line_item_description']}: GF={gf}, Other={oth}, Total={tot}{fed}")

    # Retrieve provisos
    proviso_text = ""
    try:
        voyage = get_voyage()
        idx    = get_pinecone()
        if voyage and idx:
            emb = voyage.embed(
                [f"Section {section_number} {agency_name} budget provisions"],
                model=VOYAGE_MODEL, input_type="query"
            ).embeddings[0]
            results = idx.query(
                vector=emb, top_k=10, namespace=PINECONE_NS,
                filter={"linked_section": {"$eq": section_number}},
                include_metadata=True,
            )
            if not results.matches or all(m.score < MIN_SCORE for m in results.matches):
                results = idx.query(vector=emb, top_k=6, namespace=PINECONE_NS, include_metadata=True)
            parts = [
                f"[{(m.metadata or {}).get('source_pdf','')}]\n{(m.metadata or {}).get('text_preview','')}"
                for m in results.matches if m.score >= MIN_SCORE
            ]
            if parts:
                proviso_text = "\n\nRELEVANT PROVISOS:\n\n" + "\n\n---\n\n".join(parts[:6])
    except Exception as e:
        log.warning("Pinecone failed in _get_structured_insights: %s", e)

    user_message = "\n".join(prompt_lines) + proviso_text + (
        f"\n\n---\n\nProduce the structured JSON block and full ZBB report for {agency_name} (Section {section_number})."
    )

    try:
        resp = claude.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=4096,
            system=STRUCTURED_INSIGHTS_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_message}],
        )
        raw = resp.content[0].text
    except Exception as e:
        log.warning("Claude call failed in _get_structured_insights: %s", e)
        return {"units": [], "narrative": ""}

    # Parse <structured>...</structured> block
    units: list = []
    narrative = raw
    try:
        start = raw.index("<structured>") + len("<structured>")
        end   = raw.index("</structured>")
        json_block = raw[start:end].strip()
        units = _json.loads(json_block)
        narrative = raw[end + len("</structured>"):].strip()
    except (ValueError, _json.JSONDecodeError) as e:
        log.warning("Could not parse structured block: %s", e)

    return {"units": units, "narrative": narrative}


@app.get("/agency/{section_number}/insights")
def agency_insights(section_number: str, structured: bool = False):
    """
    Claude-powered ZBB analysis for a specific agency.
    Combines actual H.4025 line items with policy knowledge and peer benchmarks.
    Returns a structured markdown report. Allow 30-60 seconds — this calls the LLM.
    """
    from collections import defaultdict

    conn = get_db()

    recap = conn.execute(
        "SELECT agency_name, total_funds, general_funds FROM recapitulation WHERE section_number=?",
        (section_number,)
    ).fetchone()
    if not recap:
        conn.close()
        raise HTTPException(404, f"Section {section_number} not found.")

    # Current-year line items only (exclude prior-year comparison rows)
    items = conn.execute("""
        SELECT subsection_name, line_item_description,
               general_funds, federal_funds, other_funds, total_funds,
               has_federal_match, federal_match_note, source_doc, page_number
        FROM line_items
        WHERE section_number=? AND is_total_row=0 AND (fund_category IS NULL OR fund_category != 'prior_year')
        ORDER BY id
    """, (section_number,)).fetchall()
    conn.close()

    if not items:
        raise HTTPException(404, f"No line items found for section {section_number}.")

    agency_name   = recap["agency_name"]
    total_cents   = recap["total_funds"] or 0
    gf_cents      = recap["general_funds"] or 0
    other_cents   = total_cents - gf_cents
    total_display = cents_to_dollars_str(total_cents)
    gf_display    = cents_to_dollars_str(gf_cents)
    other_display = cents_to_dollars_str(other_cents)

    # Structured mode: delegate to helper and return with units array
    if structured:
        result = _get_structured_insights(section_number, agency_name, items)
        return {
            "agency": agency_name,
            "section": section_number,
            "fiscal_year": FISCAL_YEAR,
            "total_funds_display": total_display,
            "general_funds_display": gf_display,
            "other_funds_display": other_display,
            "line_item_count": len(items),
            "units": result["units"],
            "analysis": result["narrative"],
            "model": CLAUDE_MODEL,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "data_note": (
                "Dollar figures are verbatim from H.4025 (Part IA). "
                "Peer benchmarks and pre_fill_text are Claude policy analysis — not SC budget figures."
            ),
        }

    claude = get_claude()
    if not claude:
        raise HTTPException(503, "ANTHROPIC_API_KEY not set.")

    # Build prompt — group by subsection
    subsections: dict = defaultdict(list)
    for it in items:
        subsections[it["subsection_name"] or "GENERAL"].append(it)

    prompt_lines = [
        f"AGENCY: {agency_name}",
        f"SECTION: {section_number}",
        f"FISCAL YEAR: {FISCAL_YEAR}",
        f"SOURCE: H.4025, SC FY{FISCAL_YEAR} Appropriations Act, ratified May 28 2025",
        "",
        "APPROPRIATIONS TOTALS (Recapitulation, Part IA):",
        f"  Total Funds:   {total_display}",
        f"  General Funds: {gf_display}",
        f"  Other Funds:   {other_display}",
        "",
        "LINE ITEMS BY PROGRAM AREA:",
    ]

    for subsec, sub_items in subsections.items():
        prompt_lines.append(f"\n--- {subsec} ---")
        for it in sub_items:
            gf  = cents_to_dollars_str(it["general_funds"] or 0)
            oth = cents_to_dollars_str(it["other_funds"] or 0)
            tot = cents_to_dollars_str(it["total_funds"] or 0)
            fed = "  ⚠ FEDERAL MATCH FLAG" if it["has_federal_match"] else ""
            prompt_lines.append(f"  {it['line_item_description']}: GF={gf}, Other={oth}, Total={tot}{fed}")
            if it["has_federal_match"] and it["federal_match_note"]:
                note = (it["federal_match_note"] or "")[:300].strip()
                prompt_lines.append(f"    [Proviso note: {note}]")

    prompt_lines.append("")

    # Retrieve provisos from Pinecone
    proviso_text = ""
    try:
        voyage = get_voyage()
        idx    = get_pinecone()
        if voyage and idx:
            query_str = f"Section {section_number} {agency_name} budget appropriations provisions"
            emb = voyage.embed([query_str], model=VOYAGE_MODEL, input_type="query").embeddings[0]
            results = idx.query(
                vector=emb, top_k=15, namespace=PINECONE_NS,
                filter={"linked_section": {"$eq": section_number}},
                include_metadata=True,
            )
            if not results.matches or all(m.score < MIN_SCORE for m in results.matches):
                results = idx.query(
                    vector=emb, top_k=10, namespace=PINECONE_NS, include_metadata=True
                )
            parts = []
            for m in results.matches:
                if m.score >= MIN_SCORE:
                    meta = m.metadata or {}
                    parts.append(
                        f"[{meta.get('source_pdf', '')}, p.{meta.get('page_number', '')}]\n"
                        f"{meta.get('text_preview', '')}"
                    )
            if parts:
                proviso_text = "RELEVANT PART IB PROVISOS:\n\n" + "\n\n---\n\n".join(parts[:10])
    except Exception as e:
        log.warning("Pinecone query failed for insights/%s: %s", section_number, e)

    user_message = "\n".join(prompt_lines)
    if proviso_text:
        user_message += "\n\n" + proviso_text
    user_message += (
        f"\n\n---\n\nPlease produce a full ZBB analyst report for {agency_name} "
        f"(Section {section_number}) per the format in your instructions."
    )

    try:
        response = claude.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=4096,
            system=INSIGHTS_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_message}],
        )
        analysis_text = response.content[0].text
    except Exception as e:
        raise HTTPException(502, f"Claude analysis failed: {e}")

    return {
        "agency": agency_name,
        "section": section_number,
        "fiscal_year": FISCAL_YEAR,
        "total_funds_display": total_display,
        "general_funds_display": gf_display,
        "other_funds_display": other_display,
        "line_item_count": len(items),
        "analysis": analysis_text,
        "model": CLAUDE_MODEL,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "data_note": (
            "Dollar figures in this analysis are verbatim from H.4025 (Part IA). "
            "Peer benchmarks and efficiency estimates are Claude policy analysis — "
            "not SC budget figures. Verify all policy recommendations against official "
            "appropriations act and agency records before acting."
        ),
    }


# ════════════════════════════════════════════════════════════════════════════
# POST /scenario
# ════════════════════════════════════════════════════════════════════════════

class ScenarioRequest(BaseModel):
    section_number: str
    cut_percentage: int = Field(..., ge=0, le=100,
                                description="Integer percentage cut (0-100). No floats.")

    @field_validator("cut_percentage")
    @classmethod
    def validate_cut(cls, v):
        if not isinstance(v, int):
            raise ValueError("cut_percentage must be an integer (0-100). No floats.")
        return v


@app.post("/scenario")
def scenario(req: ScenarioRequest):
    """
    Budget cut scenario. Integer arithmetic only. No LLM.
    All figures from SQLite with citations.
    """
    conn = get_db()

    recap = conn.execute(
        "SELECT * FROM recapitulation WHERE section_number=?", (req.section_number,)
    ).fetchone()
    if not recap:
        conn.close()
        raise HTTPException(404, f"Section {req.section_number} not found")

    items = conn.execute("""
        SELECT id, subsection_name, line_item_description,
               general_funds, other_funds, total_funds,
               has_federal_match, federal_match_note,
               source_doc, page_number
        FROM line_items
        WHERE section_number=? AND is_total_row=0
        ORDER BY id
    """, (req.section_number,)).fetchall()
    conn.close()

    cut_pct = req.cut_percentage

    line_deltas = []
    orig_total_cents  = 0
    new_total_cents   = 0
    orig_gf_cents     = 0
    new_gf_cents      = 0
    orig_other_cents  = 0
    new_other_cents   = 0
    federal_warnings  = []

    for it in items:
        orig_tf = it["total_funds"] or 0
        orig_gf = it["general_funds"] or 0
        orig_of = it["other_funds"] or 0

        # Integer arithmetic cuts — NO FLOATS
        new_tf = apply_cut_integer(orig_tf, cut_pct)
        new_gf = apply_cut_integer(orig_gf, cut_pct)
        new_of = apply_cut_integer(orig_of, cut_pct)

        delta_tf = orig_tf - new_tf
        delta_gf = orig_gf - new_gf
        delta_of = orig_of - new_of

        orig_total_cents += orig_tf
        new_total_cents  += new_tf
        orig_gf_cents    += orig_gf
        new_gf_cents     += new_gf
        orig_other_cents += orig_of
        new_other_cents  += new_of

        line_deltas.append({
            "id": it["id"],
            "subsection": it["subsection_name"],
            "description": it["line_item_description"],
            "original": {
                "total_cents": orig_tf,
                "total_display": cents_to_dollars_str(orig_tf),
                "general_funds_cents": orig_gf,
                "general_funds_display": cents_to_dollars_str(orig_gf),
                "other_funds_cents": orig_of,
                "other_funds_display": cents_to_dollars_str(orig_of),
            },
            "proposed": {
                "total_cents": new_tf,
                "total_display": cents_to_dollars_str(new_tf),
                "general_funds_cents": new_gf,
                "general_funds_display": cents_to_dollars_str(new_gf),
                "other_funds_cents": new_of,
                "other_funds_display": cents_to_dollars_str(new_of),
            },
            "delta": {
                "total_cents": delta_tf,
                "total_display": cents_to_dollars_str(delta_tf),
                "general_funds_cents": delta_gf,
                "other_funds_cents": delta_of,
            },
            "has_federal_match": bool(it["has_federal_match"]),
            "citation": {
                "source_doc": it["source_doc"],
                "page_number": it["page_number"],
                "section": f"Section {req.section_number}",
                "fiscal_year": FISCAL_YEAR,
            },
        })

        # Federal match warning
        if it["has_federal_match"] and delta_of > 0:
            federal_warnings.append({
                "line_item": it["line_item_description"],
                "other_funds_cut_cents": delta_of,
                "other_funds_cut_display": cents_to_dollars_str(delta_of),
                "warning": (
                    f"This line item is flagged for potential federal match dependency. "
                    f"A {cut_pct}% cut reduces other/federal funds by "
                    f"{cents_to_dollars_str(delta_of)}. "
                    f"Actual federal fund loss depends on the applicable match rate — "
                    f"CONFIRM MATCH RATE BEFORE ASSUMING MULTIPLIED IMPACT."
                ),
                "proviso_note": (it["federal_match_note"] or "")[:400],
                "requires_confirmation": True,
            })

    total_cut_cents = orig_total_cents - new_total_cents
    gf_cut_cents    = orig_gf_cents - new_gf_cents

    return {
        "section_number": req.section_number,
        "agency_name": recap["agency_name"],
        "fiscal_year": FISCAL_YEAR,
        "cut_percentage": cut_pct,
        "arithmetic_method": "integer_division_floor",
        "note": "All arithmetic uses integer division (floor). No floating-point math.",
        "summary": {
            "original_total_cents": orig_total_cents,
            "original_total_display": cents_to_dollars_str(orig_total_cents),
            "proposed_total_cents": new_total_cents,
            "proposed_total_display": cents_to_dollars_str(new_total_cents),
            "total_cut_cents": total_cut_cents,
            "total_cut_display": cents_to_dollars_str(total_cut_cents),
            "original_gf_cents": orig_gf_cents,
            "original_gf_display": cents_to_dollars_str(orig_gf_cents),
            "proposed_gf_cents": new_gf_cents,
            "proposed_gf_display": cents_to_dollars_str(new_gf_cents),
            "gf_cut_cents": gf_cut_cents,
            "gf_cut_display": cents_to_dollars_str(gf_cut_cents),
            "citation": {
                "source_doc": "tap1a.htm",
                "section": f"Section {req.section_number}",
                "fiscal_year": FISCAL_YEAR,
                "act": "H.4025, ratified May 28 2025",
            },
        },
        "federal_match_warnings": federal_warnings,
        "federal_match_warning_count": len(federal_warnings),
        "line_items": line_deltas,
        "line_item_count": len(line_deltas),
    }


# ════════════════════════════════════════════════════════════════════════════
# POST /sandbox/export
# ════════════════════════════════════════════════════════════════════════════

class DecisionUnit(BaseModel):
    line_item_id: int
    justified_amount_cents: int = Field(..., description="Amount in integer cents")
    justification_text: str     = Field(..., max_length=2000)
    priority_tier: str          = Field(..., pattern="^(Mandated|High|Medium|Low)$")


class ExportRequest(BaseModel):
    agency_name: str
    section_number: str
    preparer_name: Optional[str] = None
    decision_units: list[DecisionUnit]
    include_insights: bool = Field(default=False)


@app.post("/sandbox/export")
def sandbox_export(req: ExportRequest):
    """
    Generate a Word document decision package.
    Every figure is pulled from SQLite with citation.
    No figures from user input are taken as authoritative budget numbers —
    the justified_amount is the analyst's proposed figure; original amounts
    from SQLite are shown alongside for comparison.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    conn = get_db()

    # Verify agency
    recap = conn.execute(
        "SELECT * FROM recapitulation WHERE section_number=?", (req.section_number,)
    ).fetchone()
    if not recap:
        conn.close()
        raise HTTPException(404, f"Section {req.section_number} not found")

    # Load all requested line items from SQLite (authoritative source)
    line_item_ids = [du.line_item_id for du in req.decision_units]
    if not line_item_ids:
        conn.close()
        raise HTTPException(400, "No decision units provided")

    placeholders = ",".join("?" * len(line_item_ids))
    db_items = conn.execute(f"""
        SELECT id, line_item_description, subsection_name,
               general_funds, other_funds, total_funds,
               source_doc, page_number, has_federal_match
        FROM line_items
        WHERE id IN ({placeholders}) AND section_number=?
        ORDER BY id
    """, (*line_item_ids, req.section_number)).fetchall()
    conn.close()

    db_map = {row["id"]: row for row in db_items}

    # Optional: run Claude structured insights to pre-fill and append analysis
    unit_map: dict = {}   # subsection_name -> {pre_fill_text, peer_benchmark}
    insights_narrative: str = ""
    if req.include_insights:
        try:
            _conn2 = get_db()
            all_items = _conn2.execute("""
                SELECT subsection_name, line_item_description,
                       general_funds, federal_funds, other_funds, total_funds,
                       has_federal_match, federal_match_note, source_doc, page_number
                FROM line_items
                WHERE section_number=? AND is_total_row=0
                  AND (fund_category IS NULL OR fund_category != 'prior_year')
                ORDER BY id
            """, (req.section_number,)).fetchall()
            _conn2.close()

            result = _get_structured_insights(req.section_number, recap["agency_name"], all_items)
            unit_map = {u["subsection"]: u for u in result.get("units", [])}
            insights_narrative = result.get("narrative", "")
        except Exception as e:
            log.warning("Structured insights failed during export: %s", e)

    # Build Word document
    doc = Document()

    # Title
    title = doc.add_heading("Zero-Based Budget Decision Package", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Agency: {recap['agency_name']}")
    doc.add_paragraph(f"Section: {req.section_number}")
    doc.add_paragraph(f"Fiscal Year: {FISCAL_YEAR}")
    doc.add_paragraph(f"Prepared: {datetime.now().strftime('%Y-%m-%d')}")
    if req.preparer_name:
        doc.add_paragraph(f"Preparer: {req.preparer_name}")
    doc.add_paragraph(
        f"Source: H.4025, SC FY{FISCAL_YEAR} Appropriations Act, ratified May 28 2025"
    )
    doc.add_paragraph(
        "NOTE: Current funding figures are verbatim extractions from Part IA (tap1a.htm). "
        "Justified amounts are analyst proposals — not enacted appropriations."
    )

    doc.add_heading("Current Appropriations (FY2025-2026)", level=1)
    doc.add_paragraph(
        f"Total Funds (Section Recapitulation): "
        f"{cents_to_dollars_str(recap['total_funds'] or 0)}"
    )
    doc.add_paragraph(
        f"General Funds (Section Recapitulation): "
        f"{cents_to_dollars_str(recap['general_funds'] or 0)}"
    )
    doc.add_paragraph(
        f"[Citation: tap1a.htm, Recapitulation, Section {req.section_number}, "
        f"H.4025 FY{FISCAL_YEAR}]"
    )

    doc.add_heading("Decision Units", level=1)

    total_justified_cents    = 0
    total_current_cents      = 0
    federal_match_items      = []

    for du in req.decision_units:
        db_item = db_map.get(du.line_item_id)
        if not db_item:
            # Line item not found for this section — skip with note
            doc.add_paragraph(
                f"[WARNING: Line item ID {du.line_item_id} not found in "
                f"Section {req.section_number} — skipped]"
            )
            continue

        current_cents   = db_item["total_funds"] or 0
        justified_cents = du.justified_amount_cents
        delta_cents     = current_cents - justified_cents

        total_current_cents   += current_cents
        total_justified_cents += justified_cents

        doc.add_heading(
            f"{db_item['line_item_description']} [{du.priority_tier}]", level=2
        )

        # Figures table
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        def set_row(table, row_idx, label, value):
            table.rows[row_idx].cells[0].text = label
            table.rows[row_idx].cells[1].text = value

        set_row(table, 0, "Current Appropriation (FY2025-2026)",
                cents_to_dollars_str(current_cents))
        set_row(table, 1, "Justified Amount (Analyst Proposal)",
                cents_to_dollars_str(justified_cents))
        set_row(table, 2, "Variance",
                cents_to_dollars_str(delta_cents))
        set_row(table, 3, "Priority Tier", du.priority_tier)

        # Citation
        doc.add_paragraph(
            f"[Source: {db_item['source_doc']}, "
            f"Page {db_item['page_number']}, "
            f"Section {req.section_number}, "
            f"H.4025 FY{FISCAL_YEAR}]"
        )

        # Justification — use Claude pre-fill if analyst left it blank
        subsec_key = db_item["subsection_name"] or ""
        insight_unit = unit_map.get(subsec_key) or {}
        jtext = du.justification_text
        if (not jtext or jtext == "(No justification provided)") and insight_unit.get("pre_fill_text"):
            jtext = insight_unit["pre_fill_text"]
        doc.add_paragraph("Justification:")
        doc.add_paragraph(jtext)

        # Peer benchmark row (only when insights are included)
        if insight_unit.get("peer_benchmark"):
            bench_p = doc.add_paragraph()
            bench_r = bench_p.add_run(f"Peer Benchmark: {insight_unit['peer_benchmark']}")
            bench_r.italic = True
            bench_r.font.size = Pt(9)

        if db_item["has_federal_match"]:
            federal_match_items.append(db_item["line_item_description"])
            doc.add_paragraph(
                "⚠ FEDERAL MATCH FLAG: This line item may have federal matching requirements. "
                "Verify applicable match rate before finalizing this decision unit."
            )

        doc.add_paragraph("")

    # Summary
    doc.add_heading("Package Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Total Current Appropriation (Selected Items)"
    summary_table.rows[0].cells[1].text = cents_to_dollars_str(total_current_cents)
    summary_table.rows[1].cells[0].text = "Total Justified Amount"
    summary_table.rows[1].cells[1].text = cents_to_dollars_str(total_justified_cents)
    summary_table.rows[2].cells[0].text = "Total Variance"
    summary_table.rows[2].cells[1].text = cents_to_dollars_str(total_current_cents - total_justified_cents)
    summary_table.rows[3].cells[0].text = "Federal Match Flagged Items"
    summary_table.rows[3].cells[1].text = str(len(federal_match_items))

    if federal_match_items:
        doc.add_paragraph("")
        doc.add_paragraph(
            "Federal Match Warning: The following items require confirmation of "
            "applicable match rates before finalizing reductions: "
            + ", ".join(federal_match_items)
        )

    doc.add_paragraph("")
    doc.add_paragraph(
        "All current funding figures are sourced verbatim from H.4025, "
        f"SC FY{FISCAL_YEAR} Appropriations Act (Part IA, tap1a.htm), ratified May 28 2025. "
        "No figures have been estimated or approximated."
    )

    # Appendix A: Claude ZBB Analysis
    if insights_narrative:
        doc.add_page_break()
        doc.add_heading("Appendix A — Claude ZBB Analysis", level=1)
        app_meta = doc.add_paragraph(
            f"Generated by {CLAUDE_MODEL} · {datetime.now().strftime('%Y-%m-%d %H:%M UTC')} · "
            "Grounded in H.4025 line items + Part IB provisos"
        )
        app_meta.runs[0].font.size = Pt(9)
        app_meta.runs[0].italic = True
        doc.add_paragraph("")

        # Render the narrative line by line
        import re as _re
        for line in insights_narrative.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                h = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                h = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                text = stripped[2:]
                for part in _re.split(r"(\*\*[^*]+\*\*)", text):
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(9.5)
                    else:
                        r = p.add_run(part); r.font.size = Pt(9.5)
            elif stripped.startswith("---"):
                pass
            else:
                p = doc.add_paragraph()
                for part in _re.split(r"(\*\*[^*]+\*\*)", stripped):
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(10)
                    else:
                        r = p.add_run(part); r.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(3)

        doc.add_paragraph("")
        disc = doc.add_paragraph(
            "DATA NOTE: Dollar figures in this appendix are verbatim from H.4025 (Part IA). "
            "Peer benchmarks, efficiency estimates, and policy analysis are generated by Claude "
            "and are not SC budget figures. Verify all recommendations against official appropriations "
            "act and agency records before acting."
        )
        disc.runs[0].italic = True
        disc.runs[0].font.size = Pt(8)

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (
        f"ZBB_Section{req.section_number}_"
        f"{datetime.now().strftime('%Y%m%d')}.docx"
    )

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ════════════════════════════════════════════════════════════════════════════
# GET /summary
# ════════════════════════════════════════════════════════════════════════════

@app.get("/summary")
def get_summary():
    """
    Complete appropriations summary: H.4025 recurring + H.4026 nonrecurring.
    All figures from SQLite. No LLM. No arithmetic beyond what SQLite provides.
    """
    conn = get_db()

    meta = {r["key"]: r["value"] for r in conn.execute(
        "SELECT key,value FROM ingestion_meta"
    ).fetchall()}

    # Authoritative recurring total from recapitulation (verbatim from tap1a.htm)
    recap_total_cents = int(meta.get("recap_total_cents", "0"))

    # Nonrecurring: pull each fund from SQLite
    nr_rows = conn.execute("""
        SELECT line_item_description, total_funds, source_doc, page_number
        FROM line_items
        WHERE fund_category = 'nonrecurring'
        ORDER BY page_number
    """).fetchall()
    conn.close()

    surplus_cents = 0
    crf_cents     = 0
    for row in nr_rows:
        desc = (row["line_item_description"] or "").lower()
        if "surplus" in desc:
            surplus_cents = row["total_funds"]
        elif "capital reserve" in desc:
            crf_cents = row["total_funds"]

    nonrecurring_cents = surplus_cents + crf_cents
    grand_total_cents  = recap_total_cents + nonrecurring_cents

    return {
        "fiscal_year": FISCAL_YEAR,
        "recurring_total": recap_total_cents // 100,
        "recurring_total_display": cents_to_dollars_str(recap_total_cents),
        "surplus": surplus_cents // 100,
        "surplus_display": cents_to_dollars_str(surplus_cents),
        "capital_reserve_fund": crf_cents // 100,
        "capital_reserve_fund_display": cents_to_dollars_str(crf_cents),
        "nonrecurring_total": nonrecurring_cents // 100,
        "nonrecurring_total_display": cents_to_dollars_str(nonrecurring_cents),
        "grand_total": grand_total_cents // 100,
        "grand_total_display": cents_to_dollars_str(grand_total_cents),
        "source": "H.4025 + H.4026 + Summary Control Document",
        "ratified": "May 28, 2025",
        "citation": {
            "h4025": "SC FY2025-2026 Appropriations Act, Part IA (tap1a.htm)",
            "h4026": "Capital Reserve Fund, Summary Control Document Line 80",
            "surplus": "FY 2024-25 Projected Surplus, Summary Control Document Line 79",
        },
    }


# ════════════════════════════════════════════════════════════════════════════
# GET /reconciliation
# ════════════════════════════════════════════════════════════════════════════

@app.get("/reconciliation")
def get_reconciliation():
    """
    Full per-agency reconciliation: DB line-item sums vs. recapitulation totals.
    Used by the Data Integrity Badge modal.
    """
    conn = get_db()
    meta = {r["key"]: r["value"] for r in conn.execute(
        "SELECT key,value FROM ingestion_meta"
    ).fetchall()}

    agencies = conn.execute("""
        SELECT
            r.section_number,
            r.agency_name,
            r.total_funds      AS recap_total,
            r.general_funds    AS recap_gf,
            COALESCE(SUM(CASE WHEN li.is_total_row=0 AND li.fund_category='recurring' THEN li.total_funds END), 0) AS db_total,
            COALESCE(SUM(CASE WHEN li.is_total_row=0 AND li.fund_category='recurring' THEN li.general_funds END), 0) AS db_gf,
            COUNT(CASE WHEN li.is_total_row=0 AND li.fund_category='recurring' THEN 1 END) AS item_count
        FROM recapitulation r
        LEFT JOIN line_items li ON li.section_number = r.section_number
        GROUP BY r.section_number, r.agency_name, r.total_funds, r.general_funds
        ORDER BY r.total_funds DESC
    """).fetchall()
    conn.close()

    rows = []
    for a in agencies:
        delta = abs((a["recap_total"] or 0) - (a["db_total"] or 0))
        rows.append({
            "section_number": a["section_number"],
            "agency_name": a["agency_name"],
            "recap_total_cents": a["recap_total"],
            "recap_total_display": cents_to_dollars_str(a["recap_total"] or 0),
            "db_total_cents": a["db_total"],
            "db_total_display": cents_to_dollars_str(a["db_total"] or 0),
            "delta_cents": delta,
            "delta_display": cents_to_dollars_str(delta),
            "status": "PASS" if delta < 100_000 else "WARN",
            "line_item_count": a["item_count"],
        })

    recap_total = int(meta.get("recap_total_cents", "0"))
    recap_gf    = int(meta.get("recap_gf_cents", "0"))

    return {
        "status": meta.get("reconciliation_status", "unknown"),
        "run_at": meta.get("last_reconciliation"),
        "fiscal_year": FISCAL_YEAR,
        "act": "H.4025, ratified May 28 2025",
        "source_doc": "tap1a.htm",
        "summary": {
            "recap_total_display": cents_to_dollars_str(recap_total),
            "recap_gf_display": cents_to_dollars_str(recap_gf),
            "recap_total_cents": recap_total,
            "agency_count": len(rows),
            "pass_count": sum(1 for r in rows if r["status"] == "PASS"),
            "warn_count": sum(1 for r in rows if r["status"] == "WARN"),
        },
        "agencies": rows,
    }


# ════════════════════════════════════════════════════════════════════════════
# Static frontend (Phase 3)
# ════════════════════════════════════════════════════════════════════════════

FRONTEND_DIST = ROOT / "frontend" / "dist"
FRONTEND_SRC  = ROOT / "frontend"

DOCS_DIR = ROOT / "docs"

if FRONTEND_DIST.exists():
    app.mount("/assets", StaticFiles(directory=str(FRONTEND_DIST / "assets")), name="assets")

    # User guide — must be registered before the SPA catch-all
    if DOCS_DIR.exists():
        app.mount("/guide/screenshots", StaticFiles(directory=str(DOCS_DIR / "screenshots")), name="guide-screenshots")

        @app.get("/guide")
        def serve_guide():
            return FileResponse(str(DOCS_DIR / "user-guide.html"))

    @app.get("/")
    def serve_frontend():
        return FileResponse(str(FRONTEND_DIST / "index.html"))

    @app.get("/{path:path}")
    def serve_spa(path: str):
        # Serve index.html for all non-API routes (SPA routing)
        file_path = FRONTEND_DIST / path
        if file_path.exists() and file_path.is_file():
            return FileResponse(str(file_path))
        return FileResponse(str(FRONTEND_DIST / "index.html"))
else:
    @app.get("/")
    def root():
        return {
            "service": "Palmetto ZBB Suite",
            "version": "2.0.0",
            "fiscal_year": FISCAL_YEAR,
            "source": "H.4025 + H.4026, SC FY2025-2026 Appropriations Act",
            "docs": "/docs",
            "endpoints": ["/health", "/summary", "/agencies", "/agency/{section_number}",
                          "/ask", "/scenario", "/sandbox/export", "/reconciliation"],
            "note": "Frontend not built. Run `npm run build` in the frontend/ directory.",
        }


# ════════════════════════════════════════════════════════════════════════════
# Entry point
# ════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8001))
    log.info("Starting Palmetto ZBB Suite on port %d", port)
    log.info("DB: %s", DB_PATH)
    uvicorn.run("execution.api_server:app", host="0.0.0.0", port=port, reload=False)
